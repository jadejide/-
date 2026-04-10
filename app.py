import json
from datetime import datetime
from io import BytesIO
from typing import Dict, Any, List
from urllib.parse import quote

import pandas as pd
import streamlit as st
from docx import Document

try:
    import gspread
except Exception:  # pragma: no cover
    gspread = None

st.set_page_config(page_title="“模型观念”素养测量指标体系专家调查问卷", layout="wide")

st.title("“模型观念”素养测量指标体系专家调查问卷")
st.caption("德尔菲法第一轮专家咨询问卷（Streamlit 单页版）")

ADMIN_EMAIL = st.secrets.get("admin", {}).get("email", "") if hasattr(st, "secrets") else ""

st.markdown(
    """
**说明**
- 本问卷仅用于学术研究，采取匿名统计方式。
- 请结合理论认识与实践经验独立判断。
- 评分采用 5 级量表：5=非常认可，4=比较认可，3=一般，2=不太认可，1=非常不认可。
- 提交后可下载本次填写结果（JSON / Excel / Word）。
- 如已配置 Google Sheets，提交结果会自动汇总到在线总表。
- 如未配置在线总表，老师可直接下载 Excel/Word 后发送到指定邮箱。
"""
)

LIKERT = [1, 2, 3, 4, 5]

second_dimensions = [
    {
        "code": "A1",
        "name": "情境识别与数学化意识",
        "desc": "指学生能够意识到现实问题可以借助数学模型进行分析，并能从情境中发现数学问题。",
    },
    {
        "code": "A2",
        "name": "数量关系抽象与模型表达",
        "desc": "指学生能够识别情境中的关键变量、数量关系和变化规律，并用适当的数学形式表示。",
    },
    {
        "code": "A3",
        "name": "模型求解与现实解释",
        "desc": "指学生能够基于所建模型进行求解，并将数学结果还原到现实情境中解释其意义。",
    },
    {
        "code": "A4",
        "name": "模型检验与反思迁移",
        "desc": "指学生能够检验模型及结果的合理性，认识模型的条件和局限，并将已有经验迁移到新情境。",
    },
]

key_performances = [
    ("A1.1", "问题情境感知", "能从生活、社会或学习情境中注意到其中包含的数量关系、空间形式或变化现象，初步意识到该情境可以作为数学分析的对象。", "A1"),
    ("A1.2", "数学问题识别", "能从具体情境中辨认出需要解决的核心问题，发现其中蕴含的数学问题，而不是停留在情境表面的叙述。", "A1"),
    ("A2.1", "关键变量识别", "能从现实情境中识别与问题解决相关的主要数量，明确已知量、未知量和变化量，抓住建模所需的核心对象。", "A2"),
    ("A2.2", "数量关系抽象", "能从具体情境中抽取数量之间的对应关系、依赖关系或变化规律，将具体事实概括为数学关系。", "A2"),
    ("A2.3", "模型形式表达", "能根据问题特点，选择并运用适当的数学形式表达数量关系和变化规律，如方程、不等式、函数、表格、图像等。", "A2"),
    ("A2.4", "数学表征释义", "能说明模型中符号、变量、参数、式子或图像所对应的现实含义，理解数学表达与现实情境之间的联系。", "A2"),
    ("A3.1", "模型求解实施", "能依据所建立的数学模型，选择适当的方法进行运算、推理或求解，得到问题结果。", "A3"),
    ("A3.2", "结果情境解释", "能将数学求解结果放回原有现实情境中进行说明，理解结果在现实中的具体指向和实际意义。", "A3"),
    ("A3.3", "现实判断决策", "能依据模型结果对现实问题作出判断、比较、选择或提出结论，使结果真正服务于问题解决。", "A3"),
    ("A4.1", "结果合理检验", "能结合现实常识、数量范围、单位关系及边界条件，对模型结果是否合理进行判断。", "A4"),
    ("A4.2", "假设局限反思", "能意识到模型建立依赖一定的简化和假设，并能认识模型的适用条件、局限性及与现实之间的差异。", "A4"),
    ("A4.3", "模型调整改进", "能在发现模型或结果不合理时，对条件、变量、关系或表达方式作出适当调整和修正。", "A4"),
    ("A4.4", "经验迁移应用", "能将已有的模型认识、分析思路或表达经验迁移到相似的新情境中，尝试用类似方式分析新问题。", "A4"),
]

evidence_items = [
    ("A1.1.1", "A1.1", "在情境材料中圈画、标出或摘录与数量、图形、位置、变化有关的信息"),
    ("A1.1.2", "A1.1", "在“你注意到了什么”类开放作答中，写出情境中的数学关注点"),
    ("A1.2.1", "A1.2", "在作答中写出该情境需要解决的核心研究问题"),
    ("A1.2.2", "A1.2", "将生活化表述改写为可计算、可比较或可判断的数学问题"),
    ("A2.1.1", "A2.1", "在题目材料中圈画、列出与求解直接相关的数量信息"),
    ("A2.1.2", "A2.1", "在作答中写出变量设定，并说明变量含义"),
    ("A2.2.1", "A2.2", "用一句简明的话写出情境中的数量关系"),
    ("A2.2.2", "A2.2", "从题目条件中写出限制关系或依赖关系"),
    ("A2.3.1", "A2.3", "在给定多种表达方式时，选出与题目关系匹配的表达形式"),
    ("A2.3.2", "A2.3", "根据题意写出方程、不等式、函数关系式等模型表达"),
    ("A2.4.1", "A2.4", "解释式子中字母表示的现实数量"),
    ("A2.4.2", "A2.4", "解释式子中某一项、某个系数或图表要素的现实意义"),
    ("A3.1.1", "A3.1", "在作答中写出基于模型的求解过程"),
    ("A3.1.2", "A3.1", "给出模型求解结果"),
    ("A3.2.1", "A3.2", "用情境语言表述结果，而非只写纯数学答案"),
    ("A3.2.2", "A3.2", "在结果表述中写明现实单位、对象或条件"),
    ("A3.3.1", "A3.3", "根据结果作出明确判断，如“选哪种方案”“是否满足条件”"),
    ("A3.3.2", "A3.3", "给出面向问题解决的结论或建议"),
    ("A4.1.1", "A4.1", "在得到结果后给出“合理/不合理”或“符合/不符合题意”的判断"),
    ("A4.1.2", "A4.1", "给出判断合理性的依据，如常识、单位、范围或边界条件"),
    ("A4.2.1", "A4.2", "写出模型成立所依赖的条件或前提"),
    ("A4.2.2", "A4.2", "指出模型未考虑的现实因素或可能产生偏差的原因"),
    ("A4.3.1", "A4.3", "指出模型中需要修改的是条件、变量、关系还是表达方式"),
    ("A4.3.2", "A4.3", "写出修正后的模型表达或更新后的结果"),
    ("A4.4.1", "A4.4", "在新情境中写出变量设定，并说明其现实含义"),
    ("A4.4.2", "A4.4", "在新情境中写出对应的关系式或基本模型框架"),
]

kp_to_dim = {code: dim for code, _, _, dim in key_performances}
kp_name = {code: name for code, name, _, _ in key_performances}
dim_name = {d["code"]: d["name"] for d in second_dimensions}


def build_excel_bytes(result: Dict[str, Any]) -> bytes:
    basic_rows = []
    basic_info = result["basic_info"]
    for key, value in basic_info.items():
        if key == "judgement_basis":
            continue
        basic_rows.append({"字段": key, "内容": value})
    for key, value in basic_info["judgement_basis"].items():
        basic_rows.append({"字段": f"判断依据::{key}", "内容": value})

    second_rows = []
    for code, item in result["second_dimensions"].items():
        second_rows.append(
            {
                "二级维度编码": code,
                "二级维度名称": item["name"],
                "重要性": item["importance"],
                "独立性": item["independence"],
                "修改意见": item["suggestion"],
            }
        )

    kp_rows = []
    for code, item in result["key_performances"].items():
        kp_rows.append(
            {
                "关键表现编码": code,
                "关键表现名称": item["name"],
                "所属二级维度": item["dimension"],
                "适切性": item["appropriateness"],
                "一致性": item["consistency"],
                "修改意见": item["suggestion"],
            }
        )

    kp_overall_rows = [{"项目": key, "内容": value} for key, value in result["key_performance_overall"].items()]

    evidence_rows = []
    for code, item in result["evidence"].items():
        evidence_rows.append(
            {
                "证据编码": code,
                "所属关键表现": item["key_performance"],
                "证据描述": item["description"],
                "代表性": item["representative"],
                "可观测性": item["observable"],
                "修改意见": item["suggestion"],
            }
        )

    evidence_overall_rows = [{"项目": key, "内容": value} for key, value in result["evidence_overall"].items()]
    overall_rows = [{"项目": key, "内容": value} for key, value in result["overall_comments"].items()]

    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        pd.DataFrame([{"提交时间": result["submitted_at"]}]).to_excel(writer, sheet_name="提交信息", index=False)
        pd.DataFrame(basic_rows).to_excel(writer, sheet_name="基本信息", index=False)
        pd.DataFrame(second_rows).to_excel(writer, sheet_name="二级维度评分", index=False)
        pd.DataFrame(kp_rows).to_excel(writer, sheet_name="关键表现评分", index=False)
        pd.DataFrame(kp_overall_rows).to_excel(writer, sheet_name="关键表现总体意见", index=False)
        pd.DataFrame(evidence_rows).to_excel(writer, sheet_name="可观测证据评分", index=False)
        pd.DataFrame(evidence_overall_rows).to_excel(writer, sheet_name="证据总体意见", index=False)
        pd.DataFrame(overall_rows).to_excel(writer, sheet_name="总体意见", index=False)

        for _, worksheet in writer.sheets.items():
            for column_cells in worksheet.columns:
                max_length = 0
                column_letter = column_cells[0].column_letter
                for cell in column_cells:
                    value = "" if cell.value is None else str(cell.value)
                    max_length = max(max_length, len(value))
                worksheet.column_dimensions[column_letter].width = min(max(max_length + 2, 12), 40)

    output.seek(0)
    return output.getvalue()


def build_word_bytes(result: Dict[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading("“模型观念”素养测量指标体系专家调查问卷填写结果", level=1)
    doc.add_paragraph(f"提交时间：{result['submitted_at']}")

    doc.add_heading("一、专家基本信息", level=2)
    for key, value in result["basic_info"].items():
        if key == "judgement_basis":
            continue
        doc.add_paragraph(f"{key}：{value}")
    doc.add_paragraph("判断依据：")
    for key, value in result["basic_info"]["judgement_basis"].items():
        doc.add_paragraph(f"- {key}：{value}")

    doc.add_heading("二、二级维度咨询", level=2)
    for code, item in result["second_dimensions"].items():
        doc.add_paragraph(f"{code} {item['name']}")
        doc.add_paragraph(f"重要性：{item['importance']}；独立性：{item['independence']}")
        if item["suggestion"]:
            doc.add_paragraph(f"修改意见：{item['suggestion']}")

    doc.add_heading("三、关键表现咨询", level=2)
    for code, item in result["key_performances"].items():
        doc.add_paragraph(f"{code} {item['name']}（所属维度：{item['dimension']}）")
        doc.add_paragraph(f"适切性：{item['appropriateness']}；一致性：{item['consistency']}")
        if item["suggestion"]:
            doc.add_paragraph(f"修改意见：{item['suggestion']}")

    doc.add_heading("关键表现总体意见", level=3)
    for key, value in result["key_performance_overall"].items():
        doc.add_paragraph(f"{key}：{value}")

    doc.add_heading("四、可观测证据咨询", level=2)
    for code, item in result["evidence"].items():
        doc.add_paragraph(f"{code}（{item['key_performance']}）{item['description']}")
        doc.add_paragraph(f"代表性：{item['representative']}；可观测性：{item['observable']}")
        if item["suggestion"]:
            doc.add_paragraph(f"修改意见：{item['suggestion']}")

    doc.add_heading("可观测证据总体意见", level=3)
    for key, value in result["evidence_overall"].items():
        doc.add_paragraph(f"{key}：{value}")

    doc.add_heading("五、总体意见", level=2)
    for key, value in result["overall_comments"].items():
        doc.add_paragraph(f"{key}：{value}")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def flatten_result(result: Dict[str, Any]) -> Dict[str, Any]:
    row: Dict[str, Any] = {"submitted_at": result["submitted_at"]}

    for key, value in result["basic_info"].items():
        if key == "judgement_basis":
            continue
        row[f"basic__{key}"] = value

    for key, value in result["basic_info"]["judgement_basis"].items():
        row[f"basis__{key}"] = value

    for code, item in result["second_dimensions"].items():
        row[f"second__{code}__importance"] = item["importance"]
        row[f"second__{code}__independence"] = item["independence"]
        row[f"second__{code}__suggestion"] = item["suggestion"]

    for code, item in result["key_performances"].items():
        row[f"kp__{code}__appropriateness"] = item["appropriateness"]
        row[f"kp__{code}__consistency"] = item["consistency"]
        row[f"kp__{code}__suggestion"] = item["suggestion"]

    for key, value in result["key_performance_overall"].items():
        row[f"kp_overall__{key}"] = value

    for code, item in result["evidence"].items():
        row[f"evidence__{code}__representative"] = item["representative"]
        row[f"evidence__{code}__observable"] = item["observable"]
        row[f"evidence__{code}__suggestion"] = item["suggestion"]

    for key, value in result["evidence_overall"].items():
        row[f"evidence_overall__{key}"] = value

    for key, value in result["overall_comments"].items():
        row[f"overall__{key}"] = value

    return row


def get_gsheet_config() -> Dict[str, str]:
    google_sheet = st.secrets.get("google_sheet", {}) if hasattr(st, "secrets") else {}
    return {
        "spreadsheet_id": google_sheet.get("spreadsheet_id", ""),
        "worksheet_name": google_sheet.get("worksheet_name", "responses"),
    }


@st.cache_resource
def get_gspread_client():
    if gspread is None:
        raise RuntimeError("缺少 gspread 依赖。")
    if "gcp_service_account" not in st.secrets:
        raise RuntimeError("未检测到 gcp_service_account secrets 配置。")
    return gspread.service_account_from_dict(dict(st.secrets["gcp_service_account"]))


def append_result_to_gsheet(result: Dict[str, Any]) -> str:
    cfg = get_gsheet_config()
    spreadsheet_id = cfg["spreadsheet_id"].strip()
    worksheet_name = cfg["worksheet_name"].strip() or "responses"

    if not spreadsheet_id:
        raise RuntimeError("未配置 google_sheet.spreadsheet_id。")

    gc = get_gspread_client()
    sh = gc.open_by_key(spreadsheet_id)
    try:
        ws = sh.worksheet(worksheet_name)
    except gspread.WorksheetNotFound:
        ws = sh.add_worksheet(title=worksheet_name, rows=1000, cols=300)

    row = flatten_result(result)
    headers: List[str] = list(row.keys())
    values = [row.get(h, "") for h in headers]

    existing_headers = ws.row_values(1)
    if not existing_headers:
        ws.append_row(headers, value_input_option="USER_ENTERED")
        ws.append_row(values, value_input_option="USER_ENTERED")
    else:
        if existing_headers != headers:
            raise RuntimeError("Google Sheets 表头与当前问卷字段不一致。请新建工作表，或清空现有工作表后重试。")
        ws.append_row(values, value_input_option="USER_ENTERED")

    return worksheet_name


def build_mailto_link(admin_email: str, submitted_at: str) -> str:
    subject = quote(f"模型观念问卷结果 {submitted_at}")
    body = quote(
        "老师您好，\n\n已填写完成问卷。请将本页面下载的 Excel 或 Word 文件作为附件发送给我。\n\n谢谢！"
    )
    return f"mailto:{admin_email}?subject={subject}&body={body}"


sheet_cfg = get_gsheet_config()
with st.expander("管理员提示", expanded=False):
    if sheet_cfg["spreadsheet_id"] and "gcp_service_account" in st.secrets:
        st.success(f"已检测到 Google Sheets 自动汇总配置，目标工作表：{sheet_cfg['worksheet_name'] or 'responses'}")
    else:
        st.info("当前未配置 Google Sheets 自动汇总。推荐使用“导出 Excel/Word 后发邮箱”的方式收集。")
    if ADMIN_EMAIL:
        st.info(f"当前收件邮箱：{ADMIN_EMAIL}")
    else:
        st.warning("尚未配置收件邮箱。可在 Streamlit secrets 中添加 [admin] email = '你的邮箱'。")


with st.form("survey_form"):
    st.header("第一部分 专家基本情况调查")
    c1, c2 = st.columns(2)
    with c1:
        gender = st.radio("您的性别", ["男", "女"], horizontal=True)
        age = st.selectbox("您的年龄", ["40周岁以下", "41-45周岁", "46-55周岁", "56周岁以上"])
        work_years = st.selectbox("您的工作年限", ["10年以下", "11-20年", "21-30年", "31年以上"])
    with c2:
        education = st.selectbox("您的最高学历", ["本科", "硕士研究生", "博士研究生"])
        title = st.selectbox("您的职称", ["初级", "中级", "副高级", "正高级", "无"])
        familiarity = st.selectbox("您对本次调研内容的熟悉程度", ["非常不熟悉", "比较不熟悉", "一般熟悉", "比较熟悉", "非常熟悉"])

    subject_research = st.text_input("您从事的学科及研究方向")

    st.subheader("您评判本研究各指标的判断依据和影响程度")
    st.caption("请在每一行中选择影响程度")
    basis = {}
    for item in ["实践经验", "理论分析", "国内外同行了解", "直观感觉"]:
        basis[item] = st.radio(item, ["影响程度小", "影响程度中", "影响程度大"], horizontal=True, key=f"basis_{item}")

    st.header("第二部分 二级维度咨询")
    st.caption("评分说明：请从重要性、独立性两个方面进行评分，并填写修改意见。")
    second_scores = {}
    for d in second_dimensions:
        with st.expander(f"{d['code']} {d['name']}", expanded=False):
            st.write(d["desc"])
            col1, col2 = st.columns(2)
            with col1:
                importance = st.select_slider(f"{d['code']} 重要性", options=LIKERT, value=3, key=f"imp_{d['code']}")
            with col2:
                independence = st.select_slider(f"{d['code']} 独立性", options=LIKERT, value=3, key=f"ind_{d['code']}")
            suggestion = st.text_area(f"{d['code']} 修改意见", key=f"sug_{d['code']}")
            second_scores[d["code"]] = {
                "name": d["name"],
                "importance": importance,
                "independence": independence,
                "suggestion": suggestion,
            }

    st.header("第三部分 关键表现咨询")
    st.caption("评分说明：请从适切性、一致性两个方面进行评分，并填写修改意见。")
    kp_scores = {}
    for code, name, desc, dim in key_performances:
        with st.expander(f"{code} {name}（所属维度：{dim_name[dim]}）", expanded=False):
            st.write(desc)
            col1, col2 = st.columns(2)
            with col1:
                appropriateness = st.select_slider(f"{code} 适切性", options=LIKERT, value=3, key=f"app_{code}")
            with col2:
                consistency = st.select_slider(f"{code} 一致性", options=LIKERT, value=3, key=f"con_{code}")
            suggestion = st.text_area(f"{code} 修改意见", key=f"kpsug_{code}")
            kp_scores[code] = {
                "name": name,
                "dimension": dim,
                "appropriateness": appropriateness,
                "consistency": consistency,
                "suggestion": suggestion,
            }

    kp_overall = {
        "建议合并": st.text_input("哪些关键表现建议合并"),
        "建议删除": st.text_input("哪些关键表现建议删除"),
        "建议新增": st.text_input("哪些关键表现建议新增"),
        "更适合作为可观测证据": st.text_input("哪些条目更适合作为“可观测证据”而非“关键表现”"),
    }

    st.header("第四部分 可观测证据咨询")
    st.caption("评分说明：请从代表性、可观测性两个方面进行评分，并填写修改意见。")
    evidence_scores = {}
    for code, kp, desc in evidence_items:
        with st.expander(f"{code}（{kp_name[kp]}）", expanded=False):
            st.write(desc)
            st.caption(f"所属维度：{dim_name[kp_to_dim[kp]]} / 所属关键表现：{kp}")
            col1, col2 = st.columns(2)
            with col1:
                representative = st.select_slider(f"{code} 代表性", options=LIKERT, value=3, key=f"rep_{code}")
            with col2:
                observable = st.select_slider(f"{code} 可观测性", options=LIKERT, value=3, key=f"obs_{code}")
            suggestion = st.text_area(f"{code} 修改意见", key=f"evsug_{code}")
            evidence_scores[code] = {
                "key_performance": kp,
                "description": desc,
                "representative": representative,
                "observable": observable,
                "suggestion": suggestion,
            }

    ev_overall = {
        "不够直接难以观测": st.text_input("哪些证据不够直接、难以观测"),
        "与关键表现重复": st.text_input("哪些证据与关键表现重复"),
        "应删除或替换": st.text_input("哪些证据应删除或替换"),
        "建议补充": st.text_input("您建议补充的核心证据"),
    }

    st.header("第五部分 总体意见")
    overall_adv = st.text_area("您认为本指标体系目前最突出的优点是")
    overall_fix = st.text_area("您认为目前最需要修改的部分是")
    overall_focus = st.text_area("您对第二轮修订的重点建议是")
    overall_other = st.text_area("您对本指标体系的补充意见或其他宝贵意见")

    submitted = st.form_submit_button("提交问卷")

if submitted:
    result = {
        "submitted_at": datetime.now().isoformat(timespec="seconds"),
        "basic_info": {
            "gender": gender,
            "age": age,
            "work_years": work_years,
            "education": education,
            "title": title,
            "subject_research": subject_research,
            "familiarity": familiarity,
            "judgement_basis": basis,
        },
        "second_dimensions": second_scores,
        "key_performances": kp_scores,
        "key_performance_overall": kp_overall,
        "evidence": evidence_scores,
        "evidence_overall": ev_overall,
        "overall_comments": {
            "strengths": overall_adv,
            "need_revision": overall_fix,
            "next_round_focus": overall_focus,
            "other_comments": overall_other,
        },
    }

    st.success("提交成功。")

    sheet_ok = False
    try:
        worksheet_name = append_result_to_gsheet(result)
        sheet_ok = True
        st.success(f"本次填写结果已自动汇总到 Google Sheets 工作表：{worksheet_name}")
    except Exception:
        pass

    if not sheet_ok:
        if ADMIN_EMAIL:
            st.info(f"当前未启用自动汇总。请下载下方 Excel 或 Word 文件，并发送至：{ADMIN_EMAIL}")
            st.markdown(f"[点击打开邮件客户端发邮件]({build_mailto_link(ADMIN_EMAIL, result['submitted_at'])})")
        else:
            st.warning("当前未启用自动汇总，也未配置收件邮箱。建议管理员在 secrets 中添加邮箱。")

    st.subheader("提交结果预览")
    st.json(result)

    json_bytes = json.dumps(result, ensure_ascii=False, indent=2).encode("utf-8")
    excel_bytes = build_excel_bytes(result)
    word_bytes = build_word_bytes(result)

    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button(
            label="下载填写结果（JSON）",
            data=json_bytes,
            file_name="model_literacy_survey_response.json",
            mime="application/json",
        )
    with col2:
        st.download_button(
            label="下载填写结果（Excel）",
            data=excel_bytes,
            file_name="model_literacy_survey_response.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    with col3:
        st.download_button(
            label="下载填写结果（Word）",
            data=word_bytes,
            file_name="model_literacy_survey_response.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    if sheet_ok:
        st.info("你后续只需打开 Google Sheets 总表，就能看到所有老师的累计提交数据。")
    elif ADMIN_EMAIL:
        st.info("老师可先下载 Excel 或 Word，再作为附件发送到你的邮箱。这个方案不需要配置 Google Cloud。")
else:
    st.info("请填写问卷并点击页面底部的“提交问卷”。")
